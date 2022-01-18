import pathlib

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import numpy as np
import cv2
import matplotlib.pyplot as plt
from matplotlib import rcParams
import matplotlib.patches as patches

import warnings
warnings.filterwarnings("ignore")

#以下設定Times New Roman字體
#plt.rcParams["font.family"] = "Times New Roman"
#以下處理標楷體字體
#plt.rcParams["font.family"] = 'sans-serif'
#plt.rcParams['font.sans-serif'] = ['DFKai-SB', 'Times New Roman']#['Microsoft JhengHei'] 
plt.rcParams['axes.unicode_minus'] = False



def extract_data(coregistration_Error_file):
    first_group = []
    second_group = []
    with open(coregistration_Error_file) as f:
        for index, line in enumerate(f.readlines()):
            time = line.split('./TCP_TW_4/M_Ss/')[-1].split('/')[0].strip()
            rrange = line.split('range:')[-1].split('azimuth:')[0].strip()
            azimuth = line.split('azimuth:')[-1].strip()
            #print(time, rrange, azimuth)
            if index % 2 == 0:
                second_group.append([time, rrange, azimuth])
            else:
                first_group.append([time, rrange, azimuth])

    return first_group, second_group


def export_chart(coregistration_Error_file, policy_path, font):

    plt.rcParams['font.sans-serif'] = font
    
    def min_max_mean_std(data):
        return [str(round(x,2)) for x in [min(data), max(data), np.mean(data), np.std(data)]]

    first_group, second_group = extract_data(coregistration_Error_file)

    G1_range = [float(x[1]) for x in first_group]
    G2_range = [float(x[1]) for x in second_group]

    G1_azimuth = [float(x[2]) for x in first_group]
    G2_azimuth = [float(x[2]) for x in second_group]

    time_periods = [x[0] for x in first_group] 

    def correct_data(title, G1, G2, verbose=True):
        if verbose:
            print('   修正 {title}...'.format(title=title))
        for idx, err in enumerate(G2):
            if err > G1[idx] and err > 0.2:
                if verbose:
                    print('      ' + time_periods[idx] + '  {g2} ===> {g1}'.format(g2=err, g1=G1[idx]))
                G2[idx] = G1[idx]

    correct_data('Range', G1_range, G2_range)
    correct_data('Azimuth', G1_azimuth, G2_azimuth)            

    fig, axs = plt.subplots(2, 1)

    plt.setp(axs, yticks=np.arange(0, 2.2, 0.2))

    
    def plot_something(index, G1, G2, title):
        ax = axs[index]
        ax.set_title(title, fontsize=30, fontweight="bold")
        ax.plot(time_periods, G1, label="第一次配準誤差")
        ax.plot(time_periods, G2, label="第二次配準誤差")
        ax.legend(loc='upper left')
        ax.set_xticklabels(time_periods, rotation=30, ha='right', fontsize=12, fontweight="bold")
        plt.setp(ax.get_yticklabels(), fontsize=12, fontweight="bold")
        #ax.set_yticklabels([0, 0.2, 0.4, 0.6, 0.8, 1.0, 1.2, 1.4, 1.8, 2.0])
        ax.set_ylim(0, 2)
        ax.set_ylabel('誤差值(Pixel)', fontsize=28, fontweight="bold")
        ax.legend(fontsize=14)
        x_max = time_periods[-1]
        ax.set_xlim([0, x_max])
        # 加入偽表格
        chart_items = [
            ['', 'min', 'max', 'mean', 'std'],
            ['第一次配準誤差'] + min_max_mean_std(G1),
            ['第二次配準誤差'] + min_max_mean_std(G2)
        ]

        # 底色
        ax.add_patch(
         patches.Rectangle(
            (0.285, 0.64),
            0.385,
            0.32,
            transform=ax.transAxes,
            facecolor = '#F0F0F0',
            fill=True
         ) )

        # 格線
        ax.hlines([0.76, 0.86], 0.2, 0.7, transform=ax.transAxes, linewidth=1, color='white')
        ax.vlines([0.43,0.49, 0.55, 0.61], 0.61, 0.99, transform=ax.transAxes, linewidth=1, color='white')

        for row in range(3):
            for col in range(5):
                txt = chart_items[row][col]
                xpos = 0.4 + 0.06 * col
                if col == 0:
                    xpos = xpos - 0.042
                ax.text(xpos, 0.9 - 0.105 * row, txt, transform=ax.transAxes, fontsize=16, horizontalalignment='center', verticalalignment='center')

    plot_something(0, G1_range, G2_range, 'Range')
    plot_something(1, G1_azimuth, G2_azimuth, 'Azimuth')
    

    plt.subplots_adjust(hspace=0.5)
    #fig.set_size_inches(11.7, 8.3)
    fig.set_size_inches(15, 10.6)
    #plt.show()
    tmp_path = policy_path + '\\tmp'
    pathlib.Path(tmp_path).mkdir(parents=True, exist_ok=True)
    img_path = tmp_path + '\\coregistration' + font.replace(' ', '-') + '.png'
    plt.savefig(img_path)
    
    #os.system('start ' + img_path)
    return img_path




if __name__ == '__main__':

    import sys
    policy_path = sys.argv[1]
    print('=== 匹配誤差報表產生器 ===')
    print('啟動中...')
    print('正在 ' + policy_path + ' 位置下尋找 Policy 資料夾... ', end='')
    
    import glob
    policies = glob.glob(policy_path + '\Policy*')
    print('共找到 ', len(policies), ' 組資料如下')

    policies = [p.split('\\')[-1] for p in policies]
    for p in policies:
        print(p)

    print('開始處理:')

    import os
    pic_groups = [[]]
    for index, policy_dir in enumerate(policies):
        # 如果沒有 Report_3coregistration_Error 檔案的話，這筆 Policy 會被跳過

        coregistration_Error_file = policy_path + '\\' + policy_dir + '\\Report_3coregistration_Error.txt'
        if not os.path.isfile(coregistration_Error_file):
            print('找不到 Report_3coregistration_Error.txt, 跳過', policy_dir)
            continue
        
        print('開始處理', policy_dir, '...')#, end = '')

        ch_img_path, en_img_path = [export_chart(coregistration_Error_file, policy_path + '\\' + policy_dir, font=font) for font in ['DFKai-SB', 'Times New Roman']]
        ch_img, en_img = [cv2.imread(p) for p in [ch_img_path, en_img_path]]

        # 剪貼兩個表格前的中文
        en_img[173:250, 517:687] = ch_img[173:250, 517:687]
        en_img[665:734, 517:687] = ch_img[665:734, 517:687]

        # 剪貼兩個表格的 Legend
        en_img[132:192, 1193:1345] = ch_img[132:192, 1193:1345]
        en_img[621:680, 1193:1345] = ch_img[621:680, 1193:1345]

        # 移動 Pixel 標題位置並剪貼 '誤差值' 過來
        pixel_text_img = en_img[195:310, 118:158].copy()
        en_img[190:400, 100:157] = (255, 255, 255)
        en_img[165:280, 92:132] = pixel_text_img
        en_img[285:403, 94:140] = ch_img[285:403, 94:140]

        en_img[680:890, 100:157] = (255, 255, 255)
        en_img[655:770, 92:132] = pixel_text_img
        en_img[770:892, 94:140] = ch_img[770:892, 94:140]

        img_path = policy_path + '\\' + policy_dir + '\\coregistration.png'
        cv2.imwrite(img_path, en_img[40:-1,0:-1])
        #os.system('start ' + img_path)
        
        # 將輸出的圖片兩張一組分開
        if len(pic_groups[-1]) == 2:
            pic_groups.append([img_path])
        else:
            pic_groups[-1].append(img_path)

    # 產生暫存資料夾
    tmp_path = policy_path + '\\tmp'
    pathlib.Path(tmp_path).mkdir(parents=True, exist_ok=True)

    # 將圖片一張一張貼進表格
    doc_paths = []
    for doc_index, pic_group in enumerate(pic_groups):
        doc = Document('templates\\coregistration.docx')
        table = doc.tables[0]
        # 改編號
        for img_index, img_path in enumerate(pic_group):
            cell = table.rows[img_index * 2].cells[0]
            cell.paragraphs[0].runs[1].text = str( img_index + doc_index * 2 + 1)

        if doc_index > 0:
            doc.paragraphs[0].text = ''

        # 貼圖
        for img_index, img_path in enumerate(pic_group):
            cell = table.rows[img_index * 2 + 1].cells[0]
            cell._element.clear_content()
            cell.add_paragraph().add_run().add_picture(img_path, height=Cm(10))
            cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

        if len(pic_group) < 2:
            def remove_row(table, row):
                tbl = table._tbl
                tr = row._tr
                tbl.remove(tr)

            # delete empty rows
            remove_row(table, table.row[2])
            remove_row(table, table.row[3])

        doc_path = tmp_path + '\\coregistration-' + str(doc_index) + '.docx'
        doc_paths.append(doc_path)
        doc.save(doc_path)

    # 組合全部的 docx
    from utils import concatenate_docx
    output_path = policy_path + "\\doc\\coregistration.docx"
    concatenate_docx(doc_paths, output_path)
    os.system(output_path)