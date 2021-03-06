from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


import os, pathlib
import numpy as np
from numpy.polynomial.polynomial import polyfit

import matplotlib.pyplot as plt
from matplotlib import rcParams
plt.rcParams["figure.figsize"] = (6,6)

#以下設定Times New Roman字體
#plt.rcParams["font.family"] = "Times New Roman"
#以下處理標楷體字體
plt.rcParams['font.sans-serif'] = ['DFKai-SB']#['Microsoft JhengHei'] 
plt.rcParams['axes.unicode_minus'] = False



def extract_data(coherence_file):
    before = []
    after = []
    with open(coherence_file) as f:
        for index, line in enumerate(f.readlines()):
            before.append(float(line.split('\t')[-2].strip()))
            after.append(float(line.split('\t')[-1].strip()))
 
    return before, after        


def export_result(coherence_file, policy_path, index):
    
    before, after = extract_data(coherence_file)
     
    plt.scatter(before, after, facecolors='none', edgecolors='b')
    plt.xlabel("濾波前同調性", fontsize=18)
    plt.ylabel("濾波後同調性", fontsize=18)
    plt.title("濾波前後同調性比較圖 (" + str(index) + ')', fontsize=20)
    plt.xlim(0, 1)
    plt.xticks([0, 0.5, 1])
    plt.minorticks_on()
    plt.ylim(0, 1)
    plt.yticks([0, 0.5, 1])
    plt.grid(which='minor',alpha=0.3)
    plt.grid(which='major',alpha=0.6, linewidth=1.5)
    
    b, m = polyfit(before, after,1)
    ggg = [min(before), max(before)]
    plt.plot(ggg, b + m * np.array(ggg), color='black', linestyle='--', dashes=(5, 5))

    #plt.show()
    img_path = policy_path + '\\coherence.png'
    plt.savefig(img_path, dpi=100)
    plt.clf()
    return img_path



if __name__ == '__main__':

    import sys
    policy_path = sys.argv[1]
    print('=== 同調性報表產生器 ===')
    print('啟動中...')
    print('正在 ' + policy_path + ' 位置下尋找 Policy 資料夾... ', end='')

    import glob
    policies = glob.glob(policy_path + '\Policy*')
    print('共找到 ', len(policies), ' 組資料如下')

    policies = [p.split('\\')[-1] for p in policies]
    for p in policies:
        print(p)

    print('開始處理:')

    import os

    pic_groups = [[]]
    for index, policy_dir in enumerate(policies):
        # 如果沒有 ifg_coh_filt_coh_compare 檔案的話，這筆 Policy 會被跳過

        coherence_file = policy_path + '\\' + policy_dir + '\\coherence_phase\\ifg_coh_filt_coh_compare'
        if not os.path.isfile(coherence_file):
            print('找不到 ifg_coh_filt_coh_compare, 跳過', policy_dir)
            continue
        
        print('開始處理', policy_dir, '...')#, end = '')

        img_path = export_result(coherence_file, policy_path + '\\' + policy_dir, index+1)

        # 將輸出的圖片六張一組分開
        if len(pic_groups[-1]) == 6:
            pic_groups.append([img_path])
        else:
            pic_groups[-1].append(img_path)

    # 產生暫存資料夾
    tmp_path = policy_path + '\\tmp'
    pathlib.Path(tmp_path).mkdir(parents=True, exist_ok=True)

    # 將圖片一張一張貼進表格
    doc_paths = []
    for doc_index, pic_group in enumerate(pic_groups):
        doc = Document('templates\\coherence.docx')
        table = doc.tables[0]

        # 改編號
        for img_index, img_path in enumerate(pic_group):
            cell = table.rows[(img_index//2) * 2].cells[img_index % 2]
            cell.paragraphs[0].runs[1].text = str( img_index + doc_index * 6 + 1)

        if doc_index > 0:
            doc.paragraphs[0].text = ''

        # 上圖
        for img_index, img_path in enumerate(pic_group):
            cell = table.rows[(img_index//2) * 2 + 1].cells[img_index % 2]
            cell._element.clear_content()
            cell.add_paragraph().add_run().add_picture(img_path, width=Cm(6))
            cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

        if len(pic_group) < 5:
            def remove_row(table, row):
                tbl = table._tbl
                tr = row._tr
                tbl.remove(tr)

            # delete empty rows
            for row in table.rows[((len(pic_group)-1)//2) * 2 + 2:]:
                remove_row(table, row)

        if doc_index != len(pic_groups) - 1:
            doc.add_page_break()


        doc_path = tmp_path + '\\coherence-' + str(doc_index) + '.docx'
        doc_paths.append(doc_path)
        doc.save(doc_path)
        
    # 組合全部的 docx
    from utils import concatenate_docx
    output_path = policy_path + "\\doc\\coherence.docx"
    concatenate_docx(doc_paths, output_path)
    os.system('start ' + output_path)