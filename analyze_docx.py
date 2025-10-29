import docx

def analyze_docx_detailed(file_path):
    """
    Analyzes a .docx file and prints its detailed structure, including paragraphs and runs.
    """
    try:
        document = docx.Document(file_path)
        print(f"Analyzing document: {file_path}\n")

        print("--- Detailed Paragraphs and Runs ---")
        for i, para in enumerate(document.paragraphs):
            if para.text.strip():
                print(f"\nParagraph {i}: '{para.text}'")
                for j, run in enumerate(para.runs):
                    print(f"  Run {j}: '{run.text}'")

        print("\n--- Tables ---")
        for i, table in enumerate(document.tables):
            print(f"\nTable {i}:")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    print(f"  Cell ({row_idx}, {col_idx}):")
                    for para_idx, para in enumerate(cell.paragraphs):
                        print(f"    Para {para_idx}: '{para.text}'")
                        for run_idx, run in enumerate(para.runs):
                            print(f"      Run {run_idx}: '{run.text}'")


    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    analyze_docx_detailed("C:\\Users\\User\\Documents\\GitHub\\RadarDocGenerator\\templates\\baseline.docx")

