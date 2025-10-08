import pathlib
import utils
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from math import ceil
from PIL import Image
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

def set_table_borders(table):
    """
    Applies specific border styling to a table:
    - Thick outer borders
    - Thin inner borders
    """
    # Define border styles
    thin_border = {"sz": 4, "val": "single", "color": "auto"}
    thick_border = {"sz": 12, "val": "single", "color": "auto"}

    # Iterate through each cell to set borders
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # Clear existing borders
            for border in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(border)

            # Determine border type based on cell position
            is_top = (i == 0)
            is_left = (j == 0)
            is_bottom = (i == len(table.rows) - 1)
            is_right = (j == len(row.cells) - 1)

            # Top border
            top_border = thick_border if is_top else thin_border
            border_el = OxmlElement('w:top')
            for key, val in top_border.items():
                border_el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(border_el)

            # Left border
            left_border = thick_border if is_left else thin_border
            border_el = OxmlElement('w:left')
            for key, val in left_border.items():
                border_el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(border_el)

            # Bottom border
            bottom_border = thick_border if is_bottom else thin_border
            border_el = OxmlElement('w:bottom')
            for key, val in bottom_border.items():
                border_el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(border_el)

            # Right border
            right_border = thick_border if is_right else thin_border
            border_el = OxmlElement('w:right')
            for key, val in right_border.items():
                border_el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(border_el)
            
            tcPr.append(tcBorders)

def fill_cell(cell, txt):
    # This function seems to assume there is a run already.
    # A safer version would be:
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = txt
    else:
        cell.text = txt

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class Policy:

    def __init__(self, policy_dir, index, areas=None):
        self.index = index
        self.areas = areas
        self.policy_dir = policy_dir

        # 把 baseline 開出來
        self.shortbaseline = []
        with open(policy_dir + '\\postprocessing\\shortbaseline') as f:
            for line in f:
                self.shortbaseline.append(line.split('\t'))

        # 開啟對應的範例
        self.policy_length = len(self.shortbaseline)
        try:
            self.document = Document("templates\\baseline.docx")
        except Exception as e:
            raise FileNotFoundError("找不到 'templates\\baseline.docx' 範本檔案，請從 'templates' 資料夾中任選一個 .docx 檔案，並將其改名為 'baseline.docx'") from e

    def _get_image_aspect_ratio(self):
        """
        Finds the first available image and returns its aspect ratio.
        """
        for idx, row in enumerate(self.shortbaseline):
            day1, day2, distance, period = self.shortbaseline[idx]
            filename = day1 + '-' + day2 + '.tflt.filt.de'
            bmp_path = self.policy_dir + '\\postprocessing\\detrend_obs_file\\' + filename + '.bmp'
            if os.path.isfile(bmp_path):
                with Image.open(bmp_path) as img:
                    w, h = img.size
                    if h > 0:
                        return w / h
        return 1.0 # Default to square if no image is found

    def _get_layout_params(self):
        """
        Dynamically calculates the best layout for the image table. It optimizes for two criteria:
        1. Preserving the original image aspect ratio.
        2. Ensuring the last row of images is as full as possible (ideally > 50%).
        """
        image_aspect_ratio = self._get_image_aspect_ratio()
        total_images = self.policy_length

        if total_images == 0:
            return (4, [3.6, 3.6]) # Return a default if no images

        layouts = []
        TOTAL_WIDTH_CM = 14.7
        TOTAL_HEIGHT_CM = 12.0

        for cell_per_row in range(3, 8): # 3 to 7
            image_rows = ceil(total_images / cell_per_row)
            if image_rows == 0: continue

            cell_height = TOTAL_HEIGHT_CM / image_rows
            cell_width = TOTAL_WIDTH_CM / cell_per_row
            if cell_height == 0: continue

            cell_aspect_ratio = cell_width / cell_height
            aspect_ratio_diff = abs(cell_aspect_ratio - image_aspect_ratio)

            # Calculate fullness score
            items_in_last_row = total_images % cell_per_row
            if items_in_last_row == 0:
                items_in_last_row = cell_per_row
            fullness_score = items_in_last_row / cell_per_row

            layouts.append({
                'cell_per_row': cell_per_row,
                'cell_size': [cell_width, cell_height],
                'aspect_ratio_diff': aspect_ratio_diff,
                'fullness_score': fullness_score
            })

        # Prefer layouts where the last row is more than half full
        good_fullness_layouts = [l for l in layouts if l['fullness_score'] > 0.5]

        best_layout = None
        if good_fullness_layouts:
            # From the "good" layouts, choose the best.
            # Sort by aspect ratio difference (primary) and fullness (secondary, descending)
            best_layout = sorted(good_fullness_layouts, key=lambda l: (l['aspect_ratio_diff'], -l['fullness_score']))[0]
        else:
            # If no layout has a well-filled last row, fall back to the best aspect ratio
            best_layout = min(layouts, key=lambda l: l['aspect_ratio_diff'])

        return (best_layout['cell_per_row'], best_layout['cell_size'])

    def check_bmp_path(self):
        for idx, row in enumerate(self.shortbaseline):
            day1, day2, distance, period = self.shortbaseline[idx]

            # 檢查檔案名
            filename = day1 + '-' + day2 + '.tflt.filt.de'
            bmp_path = self.policy_dir + '\\postprocessing\\detrend_obs_file\\' + filename + '.bmp'
            if not os.path.isfile(bmp_path):
                return bmp_path


    def export_eps_to_png(self):

        plot = 'shortbaseline_plot.eps'
        eps_image = Image.open( self.policy_dir + '\\' + plot)
        eps_image.load(scale=3)
        self.image_path = self.policy_dir + '\\' + 'shortbaseline_plot.png'
        eps_image.save(self.image_path)

    def extract_dates(self):
        days = []
        for data_row in self.shortbaseline:
            day1, day2, distance, period = data_row
            if day1 not in days:
                days.append(day1)
            if day2 not in days:
                days.append(day2)
        days.sort()
        return days
   
    def fill_dates(self):
        cell = self.document.tables[0].rows[1].cells[1]
        dates = self.extract_dates()
        fill_cell(cell, '、'.join(dates) + '。')

    def fill_index(self):
        index_run = self.document.paragraphs[2].runs[0]
        index_run.text = '(' + str(self.index) + ')'

    def fill_areas(self):
        # 預設是生成文件後再由人手動填入
        area_run = self.document.paragraphs[3].runs[0]
        areas = self.areas
        if not areas:
            return

        for i in range(2, len(areas), 2):
            areas[i] = '\n' + areas[i]
        area_run.text = '、'.join(areas) + '。'

    def fill_image_metadata(self):
        table = self.document.tables[1]
    
        row_number = ceil(self.policy_length / 2)

        # 動態增刪表格列
        current_rows = len(table.rows) - 1
        
        template_row = None
        if current_rows > 0 and len(table.rows) > 1:
            template_row = table.rows[1]

        if row_number > current_rows:
            for _ in range(row_number - current_rows):
                new_row = table.add_row()
                if template_row:
                    # 套用列高
                    new_row.height = template_row.height
                    new_row.height_rule = template_row.height_rule
                    # 複製儲存格格式 (包含底色)
                    for i, template_cell in enumerate(template_row.cells):
                        new_cell = new_row.cells[i]
                        new_tcPr = deepcopy(template_cell._tc.get_or_add_tcPr())
                        p = new_cell._tc.get_or_add_tcPr()
                        p.getparent().replace(p, new_tcPr)

        elif row_number < current_rows:
            for i in range(current_rows - row_number):
                row_to_remove = table.rows[current_rows - i]
                row_to_remove._element.getparent().remove(row_to_remove._element)

        # 清空可能存在的舊資料
        for r_idx in range(1, len(table.rows)):
            for c_idx in range(len(table.columns)):
                fill_cell(table.rows[r_idx].cells[c_idx], "")

        # 開始填表 (左半)
        for idx in range(row_number):
            if idx >= len(self.shortbaseline): break
            day1, day2, distance, period = self.shortbaseline[idx]
            cells = table.rows[idx+1].cells
            fill_cell(cells[0], str(idx + 1)) # 填寫 NO
            fill_cell(cells[1], day1 + '-' + day2)
            fill_cell(cells[2], str(round(float(distance), 3)))
            fill_cell(cells[3], period.strip())

        # 繼續填表 (右半)
        for idx in range(row_number, len(self.shortbaseline)):
            day1, day2, distance, period = self.shortbaseline[idx]
            row_index_in_table = idx - row_number + 1
            cells = table.rows[row_index_in_table].cells
            fill_cell(cells[4], str(idx + 1)) # 填寫 NO
            fill_cell(cells[5], day1 + '-' + day2)
            fill_cell(cells[6], str(round(float(distance), 3)))
            fill_cell(cells[7], period.strip())

        # 將表格內所有文字置中
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 套用邊框樣式
        set_table_borders(table)

    def _resize_and_save_image(self, bmp_path, png_path, target_h_cm):
        """
        Opens a BMP image, resizes it based on a target physical height and a set DPI,
        and saves it as a PNG. Avoids upscaling.
        """
        with Image.open(bmp_path) as img:
            w, h = img.size
            if h == 0: return # Avoid division by zero for invalid images
            aspect_ratio = w / h

            # Define a reasonable DPI for document images
            DPI = 150
            
            # Convert target height to pixels based on DPI (1 inch = 2.54 cm)
            target_h_px = int((target_h_cm / 2.54) * DPI)
            target_w_px = int(target_h_px * aspect_ratio)

            # Resize the image only if it's larger than the target
            if w > target_w_px:
                # 使用 LANCZOS 濾波器以獲得較好的縮圖品質 (舊版 Pillow 相容)
                img_resized = img.resize((target_w_px, target_h_px), Image.LANCZOS)
                img_resized.save(png_path)
            else:
                # If the original image is smaller, just convert and save
                img.save(png_path)

    def fill_image_table(self, table_index, search_dir_leaf, file_pattern, filename_suffix_to_remove):
        import glob

        # 根據參數掃描資料夾取得所有影像
        search_path = os.path.join(self.policy_dir, 'postprocessing', search_dir_leaf)
        image_paths = glob.glob(os.path.join(search_path, file_pattern))
        image_paths.sort() # 確保順序一致

        total_images = len(image_paths)
        
        # 檢查文件是否有足夠的表格
        if table_index >= len(self.document.tables):
            print(f"警告：文件中的表格數量不足。找不到索引為 {table_index} 的表格。")
            return
            
        table = self.document.tables[table_index]

        # 如果沒有圖片，清空表格並返回
        if total_images == 0:
            for row in reversed(table.rows):
                row._element.getparent().remove(row._element)
            return

        # --- 動態計算排版 ---
        image_aspect_ratio = 1.0
        with Image.open(image_paths[0]) as img:
            w, h = img.size
            if h > 0:
                image_aspect_ratio = w / h
        
        layouts = []
        TOTAL_WIDTH_CM = 14.7
        TOTAL_HEIGHT_CM = 12.0
        for cell_per_row in range(3, 8):
            image_rows = ceil(total_images / cell_per_row)
            if image_rows == 0: continue
            cell_height = TOTAL_HEIGHT_CM / image_rows
            cell_width = TOTAL_WIDTH_CM / cell_per_row
            if cell_height == 0: continue
            cell_aspect_ratio = cell_width / cell_height
            aspect_ratio_diff = abs(cell_aspect_ratio - image_aspect_ratio)
            items_in_last_row = total_images % cell_per_row
            if items_in_last_row == 0: items_in_last_row = cell_per_row
            fullness_score = items_in_last_row / cell_per_row
            layouts.append({
                'cell_per_row': cell_per_row, 'cell_size': [cell_width, cell_height],
                'aspect_ratio_diff': aspect_ratio_diff, 'fullness_score': fullness_score
            })
        
        good_fullness_layouts = [l for l in layouts if l['fullness_score'] > 0.5]
        if good_fullness_layouts:
            best_layout = sorted(good_fullness_layouts, key=lambda l: (l['aspect_ratio_diff'], -l['fullness_score']))[0]
        else:
            best_layout = min(layouts, key=lambda l: l['aspect_ratio_diff'])
        
        cell_per_row, cell_size = best_layout['cell_per_row'], best_layout['cell_size']
        # --- 排版計算結束 ---

        # 確保表格有足夠的欄位
        current_cols = len(table.columns)
        if cell_per_row > current_cols:
            for _ in range(cell_per_row - current_cols):
                table.add_column(Cm(cell_size[0]))
        
        for col in table.columns:
            col.width = Cm(cell_size[0])
        
        num_rows_needed = ceil(total_images / cell_per_row) * 2
        current_rows = len(table.rows)

        if num_rows_needed > current_rows:
            for _ in range(num_rows_needed - current_rows):
                table.add_row()
        elif num_rows_needed < current_rows:
            for i in range(current_rows - num_rows_needed):
                row_to_remove = table.rows[current_rows - 1 - i]
                row_to_remove._element.getparent().remove(row_to_remove._element)

        tmp_path = os.path.join(self.policy_dir, 'tmp', 'converted_pngs')
        pathlib.Path(tmp_path).mkdir(parents=True, exist_ok=True)

        for idx, image_path in enumerate(image_paths):
            filename = os.path.basename(image_path)
            date_str = filename.replace(filename_suffix_to_remove, '')

            row_idx = (idx // cell_per_row) * 2
            col_idx = idx % cell_per_row

            cell = table.rows[row_idx].cells[col_idx]
            fill_cell(cell, date_str)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            png_filename = filename.replace(filename_suffix_to_remove, '.png')
            png_path = os.path.join(tmp_path, png_filename)
            target_h_cm = cell_size[1]
            self._resize_and_save_image(image_path, png_path, target_h_cm)

            cell = table.rows[row_idx + 1].cells[col_idx]
            cell._element.clear_content()
            p = cell.add_paragraph()
            p.add_run().add_picture(png_path, height=Cm(cell_size[1]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_plot(self):
        image_paragraph = self.document.paragraphs[4]
        image_paragraph.clear()
        image_paragraph.add_run().add_picture(self.image_path, width=Cm(15.53))

    def export_parital_docx(self, add_page_break=True):
        self.fill_dates()
        self.fill_index()
        self.fill_areas()
        self.fill_image_metadata()
        
        # 填入第一張圖表 (BMP)
        self.fill_image_table(
            table_index=2,
            search_dir_leaf='detrend_obs_file',
            file_pattern='*.tflt.filt.de.bmp',
            filename_suffix_to_remove='.tflt.filt.de.bmp'
        )
        
        # 填入第二張圖表 (GeoTIFF)
        self.fill_image_table(
            table_index=3,
            search_dir_leaf='detrend_obs_file',
            file_pattern='*.tflt.filt.de.geo.tif',
            filename_suffix_to_remove='.tflt.filt.de.geo.tif'
        )

        self.export_eps_to_png()
        self.add_plot()
        if add_page_break:
            self.document.add_page_break()
            
        doc_path = os.path.join(self.policy_dir, 'tmp', f'baseline-{self.index}.docx')
        self.document.save(doc_path)

        return doc_path



if __name__ == '__main__':

    import sys
    policy_path = sys.argv[1]
    print('=== 雷達影像基線資訊報表產生器 ===')
    print('啟動中...')
    print('正在 ' + policy_path + ' 位置下尋找 Policy 資料夾... ', end='')

    import glob
    policies = glob.glob(policy_path + '\\Policy*')
    print('共找到 ', len(policies), ' 組資料如下')
    for p in policies:
        print(p.split('\\')[-1])

    print('開始處理:')
    doc_paths = []
    for index, policy_dir in enumerate(policies):
        # 如果沒有 postprocessing 資料夾的話，這筆 Policy 會被跳過
        if os.path.isdir(policy_dir + '\\postprocessing'):
            print('processing', policy_dir, '...', end = '')

        # 創一個 Policy 物件
        doc_index = index+1
        policy = Policy(policy_dir=policy_dir, index=doc_index)
        
        # 輸出 docx，最後一頁不要換頁
        add_page_break = (doc_index != len(policies))
        doc_path = policy.export_parital_docx(add_page_break=add_page_break)
        print('done')
        doc_paths.append(doc_path)

    # 組合所有頁面並打開檔案
    output_path = policy_path + "\\doc\\基線.docx"
    utils.concatenate_docx(doc_paths, output_path)
    os.system('start ' + output_path)
